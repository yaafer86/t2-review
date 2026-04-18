from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from sqlalchemy.orm import Session
from sqlalchemy import select
from typing import List
from datetime import datetime
import uuid
import os

from ..core.database import get_db
from ..models.schemas import Chat, Report
from pydantic import BaseModel


router = APIRouter(prefix="/api/reports", tags=["reports"])


class ReportCreate(BaseModel):
    chat_id: str = None
    title: str
    format: str  # pdf, docx, xlsx, pptx, html, md
    content: str


def generate_report_sync(report_id: str, format: str, content: str, title: str, db: Session):
    """Synchronous report generation"""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches
        import pandas as pd
        from python_pptx import Presentation
        from python_pptx.util import Inches as PptxInches
        import markdown
        
        report_result = db.execute(select(Report).where(Report.id == report_id))
        report = report_result.scalar_one_or_none()
        
        if not report:
            return
        
        filename = f"{uuid.uuid4()}.{format}"
        filepath = f"/app/uploads/{filename}"
        
        if format == "md":
            with open(filepath, 'w') as f:
                f.write(f"# {title}\n\n{content}")
        
        elif format == "html":
            html_content = markdown.markdown(content)
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head><title>{title}</title></head>
            <body>{html_content}</body>
            </html>
            """
            with open(filepath, 'w') as f:
                f.write(full_html)
        
        elif format == "docx":
            doc = DocxDocument()
            doc.add_heading(title, 0)
            for paragraph in content.split('\n'):
                doc.add_paragraph(paragraph)
            doc.save(filepath)
        
        elif format == "xlsx":
            # Parse content as simple table
            lines = content.split('\n')
            data = [line.split(',') for line in lines if line.strip()]
            df = pd.DataFrame(data[1:], columns=data[0] if len(data) > 1 else [])
            df.to_excel(filepath, index=False)
        
        elif format == "pptx":
            prs = Presentation()
            slide = prs.slides.add_slide(prs.slide_layouts[0])
            slide.shapes.title.text = title
            
            for i, paragraph in enumerate(content.split('\n')[:5]):
                if i == 0:
                    continue
                slide = prs.slides.add_slide(prs.slide_layouts[1])
                slide.shapes.title.text = f"Section {i}"
                slide.placeholders[1].text = paragraph[:200]
            
            prs.save(filepath)
        
        elif format == "pdf":
            # For PDF, we'll create HTML first and could use weasyprint
            # For now, create HTML and note that full PDF requires additional setup
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head><title>{title}</title></head>
            <body>{markdown.markdown(content)}</body>
            </html>
            """
            # Save as HTML for now (PDF generation requires additional dependencies)
            filepath = filepath.replace('.pdf', '.html')
            filename = filename.replace('.pdf', '.html')
            with open(filepath, 'w') as f:
                f.write(html_content)
        
        # Update report record
        file_size = os.path.getsize(filepath) if os.path.exists(filepath) else 0
        
        report.file_path = filepath
        report.format = format
        report.file_size = file_size
        
        db.commit()
        
    except Exception as e:
        print(f"Error generating report: {e}")
        db.rollback()


@router.post("")
async def create_report(
    report_data: ReportCreate,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    """Create a new report"""
    report = Report(
        chat_id=report_data.chat_id if report_data.chat_id else None,
        title=report_data.title,
        format=report_data.format,
        file_path="pending"
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    
    # Generate report in background
    background_tasks.add_task(
        generate_report_sync,
        str(report.id),
        report_data.format,
        report_data.content,
        report_data.title,
        db
    )
    
    return {
        "id": str(report.id),
        "title": report.title,
        "format": report.format,
        "status": "generating"
    }


@router.get("")
async def list_reports(db: Session = Depends(get_db)):
    """List all reports"""
    result = db.execute(select(Report).order_by(Report.created_at.desc()))
    reports = result.scalars().all()
    
    return [
        {
            "id": str(r.id),
            "title": r.title,
            "format": r.format,
            "file_size": r.file_size,
            "created_at": r.created_at.isoformat()
        }
        for r in reports
    ]


@router.get("/{report_id}/download")
async def download_report(report_id: str, db: Session = Depends(get_db)):
    """Download a report"""
    from fastapi.responses import FileResponse
    
    result = db.execute(select(Report).where(Report.id == report_id))
    report = result.scalar_one_or_none()
    
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    if not os.path.exists(report.file_path):
        raise HTTPException(status_code=404, detail="Report file not found")
    
    return FileResponse(
        report.file_path,
        media_type="application/octet-stream",
        filename=f"{report.title}.{report.format}"
    )
